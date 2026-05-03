from io import BytesIO
from datetime import datetime

import pandas as pd
import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from output_quality import (
    OutputQualityContext,
    guard_generated_section_text,
    sanitize_brand_intelligence_text,
    sanitize_business_kpi_text,
    sanitize_geo_roadmap_text,
    sanitize_narrative_appendix_text,
)


# =========================================================
# Formatting helpers
# =========================================================

BLUE = "1F4E79"
LIGHT_BLUE = "EAF2F8"
LIGHT_GREEN = "E2F0D9"
LIGHT_YELLOW = "FFF2CC"
LIGHT_GREY = "F7F9FC"
BORDER = "D9E2F3"
DARK_GREY = "404040"


def possessive(name):
    name = str(name)
    if name.endswith("s"):
        return f"{name}'"
    return f"{name}'s"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color=BORDER):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()

    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)

        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_text_color(cell, color_hex):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor.from_string(color_hex)


def set_cell_bold(cell):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True


def set_cell_font_size(cell, size=8.5):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(size)


def set_cell_alignment(cell, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    for paragraph in cell.paragraphs:
        paragraph.alignment = alignment


def set_document_font(document, font_name="Arial"):
    styles = document.styles
    styles["Normal"].font.name = font_name
    styles["Normal"].font.size = Pt(10.5)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        if style_name in styles:
            styles[style_name].font.name = font_name


def add_section_heading(document, text, number=None):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(16)
    paragraph.paragraph_format.space_after = Pt(8)

    label = f"{number}. {text}" if number else text
    run = paragraph.add_run(label)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(31, 78, 121)

    return paragraph


def add_subheading(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)

    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(64, 64, 64)

    return paragraph


def add_paragraph_text(document, text, size=10.5, bold=False, color=None):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(5)

    run = paragraph.add_run(str(text))
    run.font.size = Pt(size)
    run.bold = bold

    if color:
        run.font.color.rgb = RGBColor.from_string(color)

    return paragraph


def add_markdownish_text(document, text):
    for line in str(text).splitlines():
        line = line.strip()

        if not line:
            continue

        line = line.replace(
            "Recommendation Likelihood",
            "Prompted Diagnostic Fit",
        )

        if line.startswith("#### "):
            add_subheading(document, line.removeprefix("#### ").strip())
        elif line.startswith("### "):
            add_subheading(document, line.removeprefix("### ").strip())
        else:
            add_paragraph_text(document, line)


def split_markdown_table_row(line):
    return [
        cell.strip()
        for cell in line.strip().strip("|").split("|")
    ]


def is_markdown_separator_row(cells):
    if not cells:
        return False

    return all(
        "-" in cell and not cell.replace("-", "").replace(":", "").strip()
        for cell in cells
    )


def parse_markdown_table(text):
    table_lines = [
        line.strip()
        for line in str(text).splitlines()
        if line.strip().startswith("|") and line.strip().endswith("|")
    ]

    if len(table_lines) < 2:
        return pd.DataFrame()

    headers = split_markdown_table_row(table_lines[0])
    rows = []

    for line in table_lines[1:]:
        cells = split_markdown_table_row(line)

        if is_markdown_separator_row(cells):
            continue

        if len(cells) < len(headers):
            cells.extend([""] * (len(headers) - len(cells)))

        if len(cells) > len(headers):
            cells = cells[:len(headers)]

        rows.append(dict(zip(headers, cells)))

    return pd.DataFrame(rows, columns=headers)


def add_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(3)

    run = paragraph.add_run(str(text))
    run.font.size = Pt(10.2)

    return paragraph


def add_callout_box(document, title, body, fill=LIGHT_BLUE):
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    cell = table.rows[0].cells[0]
    set_cell_shading(cell, fill)
    set_cell_border(cell, "B7C9E2")

    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_after = Pt(3)

    title_run = p1.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(11)
    title_run.font.color.rgb = RGBColor(31, 78, 121)

    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)

    body_run = p2.add_run(str(body))
    body_run.font.size = Pt(10)

    document.add_paragraph()


def add_styled_table(document, df, max_rows=15, font_size=8.5):
    if df is None or df.empty:
        add_paragraph_text(document, "No data available.")
        return

    df = df.head(max_rows).copy()

    table = document.add_table(rows=1, cols=len(df.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    header_cells = table.rows[0].cells

    for i, col in enumerate(df.columns):
        header_cells[i].text = str(col)
        set_cell_shading(header_cells[i], BLUE)
        set_cell_text_color(header_cells[i], "FFFFFF")
        set_cell_bold(header_cells[i])
        set_cell_font_size(header_cells[i], font_size)
        set_cell_alignment(header_cells[i], WD_ALIGN_PARAGRAPH.CENTER)

    for row_index, (_, row) in enumerate(df.iterrows()):
        row_cells = table.add_row().cells

        for i, col in enumerate(df.columns):
            value = row[col]
            row_cells[i].text = str(value)

            if row_index % 2 == 0:
                set_cell_shading(row_cells[i], LIGHT_GREY)

            row_cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_font_size(row_cells[i], font_size)
            set_cell_border(row_cells[i], BORDER)

    document.add_paragraph()


def add_kpi_cards(document, metrics):
    table = document.add_table(rows=1, cols=len(metrics))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, (label, value) in enumerate(metrics.items()):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_border(cell, "B7C9E2")

        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

        label_run = p1.add_run(str(label))
        label_run.bold = True
        label_run.font.size = Pt(8.5)
        label_run.font.color.rgb = RGBColor(80, 80, 80)

        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

        value_run = p2.add_run(str(value))
        value_run.bold = True
        value_run.font.size = Pt(16)
        value_run.font.color.rgb = RGBColor(31, 78, 121)

    document.add_paragraph()


# =========================================================
# Chart helpers
# =========================================================

def add_chart_image(document, image_stream, width=6.4):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(image_stream, width=Inches(width))
    document.add_paragraph()


def create_visibility_bar_chart(summary_df, brand):
    df = summary_df.copy()

    df = df.sort_values(
        by="average_visibility_score",
        ascending=False
    )

    # Show top 5 competitors plus target brand
    target_df = df[df["brand"].str.lower() == brand.lower()]
    competitor_df = df[df["brand"].str.lower() != brand.lower()].head(5)

    chart_df = pd.concat([competitor_df, target_df], ignore_index=True)

    fig, ax = plt.subplots(figsize=(8.5, 4.2))

    ax.bar(
        chart_df["brand"],
        chart_df["average_visibility_score"]
    )

    ax.set_title("Average AI Visibility Score by Brand")
    ax.set_ylabel("Average Visibility Score")
    ax.set_xlabel("Brand")
    ax.tick_params(axis="x", rotation=35)
    ax.grid(axis="y", alpha=0.25)

    plt.tight_layout()

    image_stream = BytesIO()
    fig.savefig(image_stream, format="png", dpi=200, bbox_inches="tight")
    plt.close(fig)

    image_stream.seek(0)
    return image_stream


def create_share_of_voice_chart(summary_df, brand):
    df = summary_df.copy()

    df = df.sort_values(
        by="share_of_voice_percent",
        ascending=False
    )

    target_df = df[df["brand"].str.lower() == brand.lower()]
    competitor_df = df[df["brand"].str.lower() != brand.lower()].head(5)

    chart_df = pd.concat([competitor_df, target_df], ignore_index=True)

    fig, ax = plt.subplots(figsize=(8.5, 4.2))

    ax.bar(
        chart_df["brand"],
        chart_df["share_of_voice_percent"]
    )

    ax.set_title("AI Share of Voice by Brand")
    ax.set_ylabel("Share of Voice (%)")
    ax.set_xlabel("Brand")
    ax.tick_params(axis="x", rotation=35)
    ax.grid(axis="y", alpha=0.25)

    plt.tight_layout()

    image_stream = BytesIO()
    fig.savefig(image_stream, format="png", dpi=200, bbox_inches="tight")
    plt.close(fig)

    image_stream.seek(0)
    return image_stream






# =========================================================
# Data helpers
# =========================================================
def sanitize_dataframe_text(df, context):
    if df is None or df.empty:
        return df

    sanitized_df = df.copy()

    for column in sanitized_df.columns:
        sanitized_df[column] = sanitized_df[column].map(
            lambda value: (
                sanitize_narrative_appendix_text(value, context)
                if isinstance(value, str)
                else value
            )
        )

    return sanitized_df


def get_target_metrics(summary_df, brand):
    target_row = summary_df[
        summary_df["brand"].str.lower() == brand.lower()
    ]

    if target_row.empty:
        return {
            "Total Mentions": 0,
            "Avg. Visibility": 0,
            "Prompts Visible": 0,
            "Share of Voice": "0%"
        }

    row = target_row.iloc[0]

    return {
        "Total Mentions": int(row.get("total_mentions", 0)),
        "Avg. Visibility": row.get("average_visibility_score", 0),
        "Prompts Visible": int(row.get("prompts_visible", 0)),
        "Share of Voice": f"{row.get('share_of_voice_percent', 0)}%"
    }


def get_visibility_status(total_mentions, avg_visibility, share_of_voice):
    if total_mentions == 0:
        return "currently not visible"

    if avg_visibility >= 7 or share_of_voice >= 20:
        return "showing strong visibility"

    if avg_visibility >= 3 or share_of_voice >= 10:
        return "showing moderate visibility"

    return "showing low visibility"


def get_visibility_state_noun(visibility_status):
    if visibility_status == "currently not visible":
        return "current non-visible state"

    return f"current state of {visibility_status}"


def get_ai_recall_status(visibility_status):
    if visibility_status == "currently not visible":
        return "currently not detectable"

    return visibility_status


def get_visibility_gap_sentences(
    brand,
    category,
    market,
    total_mentions,
    avg_visibility,
    share_of_voice
):
    if total_mentions == 0:
        return {
            "category_association": (
                f"AI systems currently associate the tracked {category} category with better-established competitors in {market}."
            ),
            "semantic_association": (
                "Semantic association with high-intent use cases, comparison queries, local intent, and decision-stage searches is not measurable in the tested answers."
            ),
            "third_party_signals": (
                "No tested AI answers provided measurable third-party trust or recommendation signals for the brand."
            ),
            "comparison_footprint": (
                "No visible comparison footprint was detected against higher-performing benchmark brands or providers."
            ),
            "owned_territory": (
                f"No clear AI-owned territory was detected in the {market} target market."
            ),
        }

    return {
        "category_association": (
            f"AI systems produced {total_mentions} mentions for {brand}, but competitor benchmarks show room to strengthen association with {category}."
        ),
        "semantic_association": (
            f"Semantic association should be expanded from the current {avg_visibility} average visibility score into high-intent use cases, comparison queries, local intent, and decision-stage searches."
        ),
        "third_party_signals": (
            f"Measurable third-party trust or recommendation signals should be strengthened to improve from the current {share_of_voice}% share of voice."
        ),
        "comparison_footprint": (
            f"Comparison footprint remains limited relative to higher-performing benchmark brands or providers at {avg_visibility} average visibility."
        ),
        "owned_territory": (
            f"The benchmark shows {share_of_voice}% share of voice, so the next priority is to build a clearer AI-owned territory in the {market} target market."
        ),
    }


def get_competitor_leaders(summary_df, brand):
    if summary_df is None or summary_df.empty:
        return {}

    competitors = summary_df[
        summary_df["brand"].str.lower() != brand.lower()
    ].copy()

    if competitors.empty:
        return {}

    leaders = {}

    metric_map = {
        "highest_visibility": "average_visibility_score",
        "highest_mentions": "total_mentions",
        "highest_sov": "share_of_voice_percent",
    }

    for key, column in metric_map.items():
        if column not in competitors.columns:
            continue

        values = pd.to_numeric(competitors[column], errors="coerce").fillna(0)
        max_value = values.max()

        if max_value <= 0:
            continue

        row = competitors.loc[values.idxmax()]
        leaders[key] = {
            "brand": row.get("brand", ""),
            "value": row.get(column, 0),
        }

    return leaders


def build_competitor_leader_sentence(leaders):
    if not leaders:
        return (
            "No benchmark competitor generated measurable visibility, mentions, "
            "or share of voice in this run, so this benchmark does not identify "
            "a stronger competitor leader."
        )

    highest_visibility = leaders.get("highest_visibility", {})
    highest_mentions = leaders.get("highest_mentions", {})
    highest_sov = leaders.get("highest_sov", {})

    parts = []

    if highest_visibility:
        parts.append(
            f"The highest average visibility brand is {highest_visibility['brand']} "
            f"({highest_visibility['value']} avg. visibility)"
        )

    if highest_mentions:
        parts.append(
            f"The highest mention brand is {highest_mentions['brand']} "
            f"({highest_mentions['value']} mentions)"
        )

    if highest_sov:
        parts.append(
            f"The highest share-of-voice brand is {highest_sov['brand']} "
            f"({highest_sov['value']}% SOV)"
        )

    return ". ".join(parts) + "."


def has_positive_competitor_signal(item):
    signal_fields = ["mentions", "avg_visibility", "prompts_visible", "sov"]

    for field in signal_fields:
        value = pd.to_numeric(item.get(field, 0), errors="coerce")
        if pd.notna(value) and value > 0:
            return True

    return False


def get_positive_competitors(top_competitors):
    return [
        item for item in top_competitors
        if has_positive_competitor_signal(item)
    ]


def get_top_competitors(summary_df, brand, limit=3):
    if summary_df is None or summary_df.empty:
        return []

    competitors = summary_df[
        summary_df["brand"].str.lower() != brand.lower()
    ].copy()

    competitors = competitors.sort_values(
        by="average_visibility_score",
        ascending=False
    ).head(limit)

    result = []

    for _, row in competitors.iterrows():
        result.append({
            "brand": row.get("brand", ""),
            "mentions": row.get("total_mentions", 0),
            "avg_visibility": row.get("average_visibility_score", 0),
            "prompts_visible": row.get("prompts_visible", 0),
            "sov": row.get("share_of_voice_percent", 0)
        })

    return result


def build_benchmark_df(summary_df, brand):
    columns = [
        "brand",
        "total_mentions",
        "average_visibility_score",
        "prompts_visible",
        "share_of_voice_percent"
    ]

    available = [col for col in columns if col in summary_df.columns]

    df = summary_df[available].copy()

    # Client version: show top 5 competitors plus the target brand.
    target_df = df[df["brand"].str.lower() == brand.lower()]
    competitor_df = df[df["brand"].str.lower() != brand.lower()]

    competitor_df = competitor_df.sort_values(
        by="average_visibility_score",
        ascending=False
    ).head(5)

    df = pd.concat([competitor_df, target_df], ignore_index=True)

    df = df.rename(columns={
        "brand": "Brand",
        "total_mentions": "Mentions",
        "average_visibility_score": "Avg. Visibility",
        "prompts_visible": "Prompts Visible",
        "share_of_voice_percent": "Share of Voice %"
    })

    return df


def build_winners_df(top_brands_df, max_rows=12):
    if top_brands_df is None or top_brands_df.empty:
        return None

    df = top_brands_df[[
        "prompt_category",
        "brand",
        "visibility_score"
    ]].copy()

    df = df.rename(columns={
        "prompt_category": "Query Type",
        "brand": "Winning Brand",
        "visibility_score": "Score"
    })

    return df.head(max_rows)


def create_strategy_priorities_df(brand, category, market, audience, top_competitors):
    positive_competitors = get_positive_competitors(top_competitors)

    if positive_competitors:
        primary = positive_competitors[0]["brand"]
        secondary = (
            positive_competitors[1]["brand"]
            if len(positive_competitors) > 1
            else primary
        )
    else:
        primary = "No measurable competitor leader"
        secondary = "Tracked competitors not measurably visible"

    return pd.DataFrame([
        {
            "Level": "High",
            "Priority": "Category Authority",
            "Target Query Territory": f"{category} recommendations for {audience}",
            "Competitor Focus": primary,
            "Recommended Action": f"Create AI-citable guides positioning {brand} around category expertise, decision criteria, and buyer fit."
        },
        {
            "Level": "High",
            "Priority": "Use-Case Association",
            "Target Query Territory": f"{category} for high-intent use cases",
            "Competitor Focus": primary,
            "Recommended Action": f"Publish educational content explaining {possessive(brand)} relevance to common buyer needs and use-case-specific searches."
        },
        {
            "Level": "High",
            "Priority": "Decision-Stage Visibility",
            "Target Query Territory": f"decision-stage {category} searches",
            "Competitor Focus": secondary,
            "Recommended Action": "Develop decision-stage pages and recommendation content for buyers, evaluators, and decision-makers."
        },
        {
            "Level": "Medium",
            "Priority": "Comparison Visibility",
            "Target Query Territory": f"{brand} vs benchmark brands or providers",
            "Competitor Focus": primary,
            "Recommended Action": "Build fair comparison pages that explain fit, use cases, and differentiation without attacking competitors."
        },
        {
            "Level": "Medium",
            "Priority": "Market Relevance",
            "Target Query Territory": f"{category} brands or providers in {market}",
            "Competitor Focus": secondary,
            "Recommended Action": "Create market-specific content with local context, buyer concerns, and recommendation language."
        }
    ])


def create_roadmap_df(brand, category, top_competitors, metrics):
    positive_competitors = get_positive_competitors(top_competitors)

    if positive_competitors:
        primary = positive_competitors[0]["brand"]
        secondary = (
            positive_competitors[1]["brand"]
            if len(positive_competitors) > 1
            else primary
        )
        tertiary = secondary
        second_phase_focus = primary
    else:
        primary = "Category baseline"
        secondary = "Market visibility baseline"
        tertiary = "Tracked competitor set"
        second_phase_focus = secondary

    current_mentions = metrics["Total Mentions"]
    current_sov = metrics["Share of Voice"]

    return pd.DataFrame([
        {
            "Phase": "30 Days",
            "Action": f"Publish core content for {category} use cases, buyer questions, and decision criteria.",
            "Target Metric": f"Increase from {current_mentions} current mentions to more AI-detectable brand mentions.",
            "Competitor Focus": primary
        },
        {
            "Phase": "60 Days",
            "Action": "Launch comparison pages and category authority explainers.",
            "Target Metric": "Improve average visibility score and prompts visible.",
            "Competitor Focus": second_phase_focus
        },
        {
            "Phase": "90 Days",
            "Action": "Build third-party mentions, reviews, and AI-citable references.",
            "Target Metric": f"Increase share of voice from the current {current_sov} toward stronger measurable inclusion.",
            "Competitor Focus": tertiary
        }
    ])


# =========================================================
# Report sections
# =========================================================

def add_cover_page(document, brand, market, report_date):
    for _ in range(5):
        document.add_paragraph()

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title_run = title.add_run(f"{brand} {market}\nAI Visibility Report")
    title_run.bold = True
    title_run.font.size = Pt(28)
    title_run.font.color.rgb = RGBColor(31, 78, 121)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle_run = subtitle.add_run("Generative Engine Optimization Audit")
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = RGBColor(90, 90, 90)

    document.add_paragraph()

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta_run = meta.add_run(f"Market: {market} | Report Date: {report_date}")
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(90, 90, 90)

    document.add_page_break()


def add_report_overview(
    document,
    brand,
    market,
    category,
    audience,
    report_date,
    run_mode="Full Report Mode",
    prompt_limit=None
):
    add_section_heading(document, "Report Overview", "1")

    deliverable_status = "Client-deliverable full report"
    if run_mode == "Quick Test Mode":
        deliverable_status = (
            "Development-only limited-prompt output. Not client-deliverable."
        )

    overview_rows = [
        {"Field": "Target Brand", "Value": brand},
        {"Field": "Market", "Value": market},
        {"Field": "Category", "Value": category},
        {"Field": "Audience", "Value": audience},
        {"Field": "Report Date", "Value": report_date},
        {"Field": "Report Type", "Value": "AI Visibility / GEO Audit"},
        {"Field": "Run Mode", "Value": run_mode},
        {"Field": "Deliverable Status", "Value": deliverable_status},
    ]

    if run_mode == "Quick Test Mode":
        overview_rows.append({"Field": "Prompt Limit", "Value": prompt_limit})

    overview_df = pd.DataFrame([
        row for row in overview_rows
    ])

    add_styled_table(document, overview_df, max_rows=10, font_size=9.5)


def add_quick_test_warning(document, run_mode, prompt_limit=None):
    if run_mode != "Quick Test Mode":
        return

    prompt_word = "prompt" if prompt_limit == 1 else "prompts"

    add_callout_box(
        document,
        "TEST VERSION ONLY",
        (
            "Quick Test Mode - Not Client Deliverable. "
            f"This report used only {prompt_limit} {prompt_word} and is intended for development and QA only. "
            "Do not use it as a full client-facing AI visibility report."
        ),
        fill=LIGHT_YELLOW
    )


def add_executive_summary(document, brand, category, market, metrics, competitor_leaders):
    add_section_heading(document, "Executive Summary", "2")

    total_mentions = metrics["Total Mentions"]
    avg_visibility = metrics["Avg. Visibility"]
    prompts_visible = metrics["Prompts Visible"]
    share_of_voice = float(str(metrics["Share of Voice"]).replace("%", "") or 0)
    visibility_status = get_visibility_status(
        total_mentions,
        avg_visibility,
        share_of_voice
    )
    ai_recall_status = get_ai_recall_status(visibility_status)
    gap_sentences = get_visibility_gap_sentences(
        brand,
        category,
        market,
        total_mentions,
        avg_visibility,
        share_of_voice
    )
    top_competitor_sentence = build_competitor_leader_sentence(competitor_leaders)

    core_finding = (
        f"{brand} is {visibility_status} in AI-generated {category} category recommendations in {market}. "
        f"The brand records {total_mentions} total mentions, an average visibility score of "
        f"{avg_visibility}, {prompts_visible} prompts visible, and "
        f"{metrics['Share of Voice']} share of voice. {top_competitor_sentence}"
    )

    add_callout_box(document, "Core Finding", core_finding, fill=LIGHT_BLUE)
    add_kpi_cards(document, metrics)

    add_subheading(document, "Key Findings")

    findings = [
        f"{brand} is {visibility_status} in the tested prompt set, with {total_mentions} total mentions and {metrics['Share of Voice']} share of voice.",
        gap_sentences["category_association"],
        "The highest-value opportunity is to connect the brand with high-intent use cases, comparison queries, local intent, and decision-stage searches.",
        f"The immediate GEO objective is to create AI-citable evidence that can improve from the current {avg_visibility} average visibility score."
    ]

    for item in findings:
        add_bullet(document, item)

def add_visual_benchmark(document, summary_df, brand):
    add_section_heading(document, "Visual Benchmark", "3")

    add_paragraph_text(
        document,
        "The charts below summarize overall AI visibility and share of voice among the most visible benchmark competitors."
    )

    visibility_chart = create_visibility_bar_chart(summary_df, brand)
    add_chart_image(document, visibility_chart, width=6.3)

    share_chart = create_share_of_voice_chart(summary_df, brand)
    add_chart_image(document, share_chart, width=6.3)



def add_competitive_benchmark(document, benchmark_df):
    add_section_heading(document, "Competitive Benchmark", "4")

    add_paragraph_text(
        document,
        "The table below summarizes the most relevant benchmark competitors plus the target brand."
    )

    add_styled_table(document, benchmark_df, max_rows=8, font_size=8.5)


def add_top_winners(document, winners_df):
    add_section_heading(document, "Top Brand Winners by Query Type", "5")

    add_paragraph_text(
    document,
    "This table shows which brand wins specific query contexts based on visibility score. "
    "It should be read together with the Competitive Benchmark section: the benchmark summarizes overall visibility, "
    "while this table shows category-level winners."
)

    if winners_df is None or winners_df.empty:
        add_paragraph_text(document, "No positive brand winners were detected.")
    else:
        add_styled_table(document, winners_df, max_rows=12, font_size=8)

def add_gap_diagnosis(document, brand, category, market, metrics, top_competitors):
    add_section_heading(document, "Visibility Gap Diagnosis", "6")
    total_mentions = metrics["Total Mentions"]
    avg_visibility = metrics["Avg. Visibility"]
    share_of_voice = float(str(metrics["Share of Voice"]).replace("%", "") or 0)
    visibility_status = get_visibility_status(
        total_mentions,
        avg_visibility,
        share_of_voice
    )
    ai_recall_status = get_ai_recall_status(visibility_status)
    gap_sentences = get_visibility_gap_sentences(
        brand,
        category,
        market,
        total_mentions,
        avg_visibility,
        share_of_voice
    )

    add_callout_box(
        document,
        "Diagnosis",
        f"{brand} is {visibility_status} in the AI recommendation set. The brand has "
        f"{total_mentions} mentions, {avg_visibility} average visibility, "
        f"and {metrics['Share of Voice']} share of voice.",
        fill=LIGHT_YELLOW
    )

    add_subheading(document, "Key Barriers")

    barriers = [
        f"Measured AI recall is {ai_recall_status}, based on {total_mentions} mentions across the tested prompt set.",
        gap_sentences["semantic_association"],
        gap_sentences["third_party_signals"],
        gap_sentences["comparison_footprint"],
        gap_sentences["owned_territory"]
    ]

    for item in barriers:
        add_bullet(document, item)

    if top_competitors:
        add_subheading(document, "Competitive Context")

        positive_competitors = [
            item for item in top_competitors
            if has_positive_competitor_signal(item)
        ]

        if not positive_competitors:
            add_bullet(
                document,
                "No benchmark competitor generated measurable AI recall signal in this run. "
                "Competitors should be treated as unproven in this dataset, not as stronger AI visibility leaders."
            )
        else:
            for item in positive_competitors:
                add_bullet(
                    document,
                    f"{item['brand']} has {item['mentions']} mentions and {item['sov']}% share of voice, creating a stronger AI recall signal than {brand}."
                )


def add_strategy_priorities(document, priorities_df):
    add_section_heading(document, "Strategic Priorities", "7")

    add_paragraph_text(
        document,
        "The following priorities translate the visibility gap into concrete GEO actions."
    )

    add_styled_table(document, priorities_df, max_rows=10, font_size=7.3)


def add_brand_intelligence(
    document,
    brand_intelligence,
    section_number=None,
    heading_text="Brand Intelligence & Positioning Audit"
):
    if not brand_intelligence:
        return

    add_section_heading(
        document,
        heading_text,
        str(section_number) if section_number else None
    )
    add_callout_box(
        document,
        "Diagnostic Insight",
        "Diagnostic insight. Tracked competitors are included in visibility scoring and share of voice. "
        "AI-discovered market signals are diagnostic references only and are not included in scoring unless selected as tracked competitors before the benchmark run.",
        fill=LIGHT_BLUE
    )

    add_subheading(document, "Recommendation Drivers")
    add_markdownish_text(
        document,
        brand_intelligence.get("recommendation_drivers", "")
    )

    add_subheading(document, "AI-Inferred Target Brand Understanding")
    add_markdownish_text(
        document,
        brand_intelligence.get("target_brand_understanding", "")
    )

    add_subheading(document, "Positioning Gap Analysis")
    add_markdownish_text(
        document,
        brand_intelligence.get("positioning_gap_analysis", "")
    )


def add_geo_content_roadmap(document, geo_content_roadmap, section_number):
    if not geo_content_roadmap:
        return

    add_section_heading(
        document,
        "GEO Content Roadmap",
        str(section_number)
    )
    add_callout_box(
        document,
        "Strategic Execution Plan",
        "Strategic execution plan. Not part of visibility scoring or share of voice.",
        fill=LIGHT_GREEN
    )

    roadmap_df = parse_markdown_table(geo_content_roadmap)
    if not roadmap_df.empty:
        add_styled_table(document, roadmap_df, max_rows=10, font_size=6.4)
    else:
        add_markdownish_text(document, geo_content_roadmap)


def add_roadmap(document, roadmap_df, section_number="8"):
    add_section_heading(document, "30 / 60 / 90 Day Roadmap", section_number)

    add_styled_table(document, roadmap_df, max_rows=10, font_size=8)


def add_measurement_plan(
    document,
    brand,
    metrics,
    section_number="9",
    quality_context=None,
):
    add_section_heading(document, "Measurement Plan", section_number)

    add_paragraph_text(
        document,
        "The next benchmark should evaluate whether the visibility gap is beginning to close."
    )

    measurement_df = pd.DataFrame([
        {
            "Metric": "Total Mentions",
            "Current State": str(metrics["Total Mentions"]),
            "Next Benchmark Target": "Begin generating detectable mentions in a future full benchmark."
        },
        {
            "Metric": "Average Visibility Score",
            "Current State": str(metrics["Avg. Visibility"]),
            "Next Benchmark Target": "Begin improving average visibility score in a future full benchmark."
        },
        {
            "Metric": "Prompts Visible",
            "Current State": str(metrics["Prompts Visible"]),
            "Next Benchmark Target": "Begin generating prompt-level visibility in a future full benchmark."
        },
        {
            "Metric": "Share of Voice",
            "Current State": metrics["Share of Voice"],
            "Next Benchmark Target": "Begin generating measurable share of voice in a future full benchmark."
        }
    ])

    if quality_context is not None:
        measurement_df = sanitize_dataframe_text(measurement_df, quality_context)

    add_styled_table(document, measurement_df, max_rows=10, font_size=8)


def add_recommended_next_step(
    document,
    brand,
    category,
    market,
    metrics,
    section_number="10"
):
    add_section_heading(document, "Recommended Next Step", section_number)
    total_mentions = metrics["Total Mentions"]
    avg_visibility = metrics["Avg. Visibility"]
    share_of_voice = float(str(metrics["Share of Voice"]).replace("%", "") or 0)
    visibility_status = get_visibility_status(
        total_mentions,
        avg_visibility,
        share_of_voice
    )
    visibility_state_noun = get_visibility_state_noun(visibility_status)

    add_callout_box(
        document,
        "Immediate Action",
        f"Build AI-citable content that connects {brand} with high-intent use cases, comparison queries, local intent, "
        f"decision-stage searches, and market-specific category queries for {category} in {market}. The next benchmark should track whether the brand "
        f"improves from its {visibility_state_noun} toward stronger inclusion in AI-generated recommendation lists.",
        fill=LIGHT_GREEN
    )

def normalize_prompt_categories(prompt_categories):
    if not prompt_categories:
        return []

    categories = []
    seen = set()

    for item in prompt_categories:
        category = str(item).strip()
        key = category.lower()

        if category and key not in seen:
            categories.append(category)
            seen.add(key)

    return categories


def add_methodology_notes(
    document,
    category,
    section_number="11",
    prompt_categories=None
):
    add_section_heading(document, "Methodology Notes", section_number)

    notes = [
    f"The benchmark is based on fixed and AI-generated prompts designed to simulate {category} recommendation queries.",
    "Visibility is calculated from brand mentions, estimated ranking, and prompt-level appearance.",
    "Share of voice reflects the distribution of brand mentions among tracked competitors.",
    "Scores reflect AI answer visibility, not actual revenue, market share, product performance, customer satisfaction, or business outcomes.",
    "The output should be interpreted as an AI visibility benchmark, not as a consumer survey, sales performance report, or clinical evaluation.",
    "Results should be re-run periodically to track whether content and visibility interventions improve AI recall."
]

    for note in notes:
        add_bullet(document, note)

    prompt_categories = normalize_prompt_categories(prompt_categories)
    if prompt_categories:
        add_paragraph_text(
            document,
            "Query intent coverage included:",
            bold=True
        )

        for prompt_category in prompt_categories:
            add_bullet(document, prompt_category)


# =========================================================
# Main function called by app.py
# =========================================================

def create_executive_docx_report(
    brand,
    market,
    category,
    audience,
    summary_df,
    top_brands_df,
    recommendations,
    strategy_report,
    gap_analysis,
    run_mode="Full Report Mode",
    prompt_limit=None,
    brand_intelligence=None,
    prompt_categories=None,
    geo_content_roadmap=None
):
    document = Document()

    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    set_document_font(document, "Arial")

    report_date = datetime.now().strftime("%Y-%m-%d")

    metrics = get_target_metrics(summary_df, brand)
    tracked_competitors = []
    if summary_df is not None and "brand" in summary_df.columns:
        tracked_competitors = [
            str(item)
            for item in summary_df["brand"].dropna().tolist()
            if str(item).strip().lower() != str(brand).strip().lower()
        ]
    quality_context = OutputQualityContext(
        category=category,
        run_mode=run_mode,
        brand=brand,
        market=market,
        audience=audience,
        tracked_competitors=tracked_competitors,
        target_brand_metrics={
            "total_mentions": metrics["Total Mentions"],
            "prompts_visible": metrics["Prompts Visible"],
            "share_of_voice_percent": float(str(metrics["Share of Voice"]).replace("%", "") or 0),
        },
    )
    if brand_intelligence:
        brand_intelligence = {
            key: (
                sanitize_brand_intelligence_text(
                    guard_generated_section_text(
                        value,
                        quality_context,
                        f"Brand Intelligence {key}",
                    ),
                    quality_context,
                )
                if isinstance(value, str)
                else value
            )
            for key, value in brand_intelligence.items()
        }
    if geo_content_roadmap:
        geo_content_roadmap = guard_generated_section_text(
            geo_content_roadmap,
            quality_context,
            "GEO Content Roadmap",
        )
        geo_content_roadmap = sanitize_geo_roadmap_text(
            geo_content_roadmap,
            quality_context,
        )
    strategy_report = guard_generated_section_text(
        strategy_report,
        quality_context,
        "AI Visibility Strategy Deep Dive",
    )
    strategy_report = sanitize_narrative_appendix_text(
        strategy_report,
        quality_context,
    )
    gap_analysis = guard_generated_section_text(
        gap_analysis,
        quality_context,
        "Gap Analysis",
    )
    gap_analysis = sanitize_narrative_appendix_text(
        gap_analysis,
        quality_context,
    )
    top_competitors = get_top_competitors(summary_df, brand, limit=3)
    competitor_leaders = get_competitor_leaders(summary_df, brand)
    benchmark_df = build_benchmark_df(summary_df, brand)
    winners_df = build_winners_df(top_brands_df, max_rows=12)
    priorities_df = create_strategy_priorities_df(
    brand,
    category,
    market,
    audience,
    top_competitors
)
    priorities_df = sanitize_dataframe_text(priorities_df, quality_context)

    roadmap_df = create_roadmap_df(brand, category, top_competitors, metrics)
    roadmap_df = sanitize_dataframe_text(roadmap_df, quality_context)

    add_cover_page(document, brand, market, report_date)
    add_report_overview(
        document,
        brand,
        market,
        category,
        audience,
        report_date,
        run_mode=run_mode,
        prompt_limit=prompt_limit
    )
    add_quick_test_warning(document, run_mode, prompt_limit)
    add_executive_summary(
        document,
        brand,
        category,
        market,
        metrics,
        competitor_leaders
    )
    add_visual_benchmark(document, summary_df, brand)
    add_competitive_benchmark(document, benchmark_df)
    add_top_winners(document, winners_df)
    add_gap_diagnosis(document, brand, category, market, metrics, top_competitors)
    add_strategy_priorities(document, priorities_df)

    next_section_number = 8

    if geo_content_roadmap:
        add_geo_content_roadmap(document, geo_content_roadmap, next_section_number)
        next_section_number += 1

    add_roadmap(document, roadmap_df, str(next_section_number))
    next_section_number += 1
    add_measurement_plan(
    document,
    brand,
    metrics,
    str(next_section_number),
    quality_context=quality_context,
)
    next_section_number += 1
    add_methodology_notes(
        document,
        category,
        str(next_section_number),
        prompt_categories=prompt_categories
    )
    if brand_intelligence:
        add_brand_intelligence(
            document,
            brand_intelligence,
            heading_text="Appendix: Brand Intelligence & Positioning Audit"
        )

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    return buffer.getvalue()
