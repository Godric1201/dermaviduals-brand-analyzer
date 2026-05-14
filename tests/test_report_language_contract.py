import re
from io import BytesIO

import pandas as pd
from docx import Document

from geo_audit.markdown_report import build_executive_markdown_report
from geo_audit.report_generator import create_executive_docx_report


FORBIDDEN_OVERCLAIM_PATTERNS = [
    r"\bguarantees\b",
    r"\bguaranteed\b",
    r"\bwill\s+increase\s+sales\b",
    r"\bwill\s+drive\s+revenue\b",
    r"\bwill\s+improve\s+rankings\b",
    r"\bwill\s+increase\s+share\s+of\s+voice\b",
    r"\bconsumers\s+prefer\b",
    r"\bmarket\s+leader\b",
    r"\bdominates\s+the\s+market\b",
    r"\bAI\s+knows\b",
    r"\bAI\s+trusts\b",
    r"\bcaused\s+retrieval\b",
    r"\bthe\s+model\s+used\s+this\s+source\b",
]


def assert_has_phrase_family(text, phrases):
    normalized = str(text).lower()

    assert any(phrase.lower() in normalized for phrase in phrases)


def assert_forbidden_overclaims_absent(text):
    for pattern in FORBIDDEN_OVERCLAIM_PATTERNS:
        assert not re.search(pattern, str(text), flags=re.IGNORECASE)


def sample_summary_df():
    return pd.DataFrame([
        {
            "brand": "Espresso House",
            "total_mentions": 2,
            "average_visibility_score": 3.5,
            "prompts_visible": 2,
            "share_of_voice_percent": 20.0,
        },
        {
            "brand": "Coffee Fellows",
            "total_mentions": 8,
            "average_visibility_score": 7.2,
            "prompts_visible": 4,
            "share_of_voice_percent": 80.0,
        },
    ])


def sample_top_brands_df():
    return pd.DataFrame([
        {
            "prompt_category": "Best Options",
            "brand": "Coffee Fellows",
            "visibility_score": 7.0,
        },
    ])

def zero_visibility_summary_df():
    return pd.DataFrame([
        {
            "brand": "Regional Re",
            "total_mentions": 0,
            "average_visibility_score": 0,
            "prompts_visible": 0,
            "share_of_voice_percent": 0,
        },
        {
            "brand": "Munich Re",
            "total_mentions": 18,
            "average_visibility_score": 75.0,
            "prompts_visible": 6,
            "share_of_voice_percent": 45.0,
        },
        {
            "brand": "Swiss Re",
            "total_mentions": 14,
            "average_visibility_score": 68.0,
            "prompts_visible": 5,
            "share_of_voice_percent": 35.0,
        },
    ])


def zero_visibility_top_brands_df():
    return pd.DataFrame([
        {
            "prompt_category": "Best Options",
            "brand": "Munich Re",
            "visibility_score": 75.0,
        },
        {
            "prompt_category": "Local Recommendations",
            "brand": "Swiss Re",
            "visibility_score": 60.0,
        },
    ])


def zero_visibility_markdown_report(**overrides):
    summary_df = zero_visibility_summary_df()
    inputs = {
        "brand": "Regional Re",
        "display_brand": "Regional Re",
        "category": "reinsurance",
        "display_category": "Reinsurance",
        "market": "Taiwan and Asia-Pacific",
        "display_market": "Taiwan and Asia-Pacific",
        "audience": "enterprise insurance buyers",
        "display_audience": "Enterprise Insurance Buyers",
        "run_mode": "Full Report Mode",
        "prompt_limit": None,
        "deliverable_status": "Client-deliverable full report",
        "summary_df": summary_df,
        "summary_display_df": summary_df.copy(),
        "detailed_pivot_df": pd.DataFrame([
            {
                "prompt_category": "Best Options",
                "Regional Re": 0,
                "Munich Re": 75.0,
                "Swiss Re": 68.0,
            },
            {
                "prompt_category": "Local Recommendations",
                "Regional Re": 0,
                "Munich Re": 65.0,
                "Swiss Re": 60.0,
            },
        ]),
        "top_brands_df": zero_visibility_top_brands_df(),
        "recommendations": "Prioritize first-detection evidence assets.",
        "plan": "Use future benchmark cycles to validate candidate-set inclusion.",
        "gap_analysis": "Regional Re was not retrieved in the tested prompt set.",
        "brand_win_explanation": None,
        "replacement_strategy": None,
        "brand_intelligence": None,
        "brand_intelligence_done": False,
        "geo_content_roadmap": None,
        "geo_content_roadmap_done": False,
        "prompt_categories": ["Best Options", "Local Recommendations"],
        "tracked_competitors": ["Munich Re", "Swiss Re"],
    }
    inputs.update(overrides)

    return build_executive_markdown_report(**inputs)

def markdown_report(**overrides):
    summary_df = sample_summary_df()
    inputs = {
        "brand": "Espresso House",
        "display_brand": "Espresso House",
        "category": "cafes",
        "display_category": "Cafes",
        "market": "Berlin",
        "display_market": "Berlin",
        "audience": "remote workers",
        "display_audience": "Remote Workers",
        "run_mode": "Full Report Mode",
        "prompt_limit": None,
        "deliverable_status": "Client-deliverable full report",
        "summary_df": summary_df,
        "summary_display_df": summary_df.copy(),
        "detailed_pivot_df": pd.DataFrame([
            {
                "prompt_category": "Best Options",
                "Espresso House": 2.0,
                "Coffee Fellows": 7.0,
            },
        ]),
        "top_brands_df": sample_top_brands_df(),
        "recommendations": "Prioritize AI-citable category evidence.",
        "plan": "Use future benchmark cycles to validate progress.",
        "gap_analysis": "The tested answers suggest limited prompt-level visibility.",
        "brand_win_explanation": None,
        "replacement_strategy": None,
        "brand_intelligence": {
            "recommendation_drivers": "Diagnostic inference from benchmark answers.",
            "target_brand_understanding": "AI-inferred target brand understanding.",
            "positioning_gap_analysis": "Positioning gaps require validation.",
        },
        "brand_intelligence_done": True,
        "geo_content_roadmap": None,
        "geo_content_roadmap_done": False,
        "prompt_categories": ["Best Options", "Local Recommendations"],
        "tracked_competitors": ["Coffee Fellows"],
    }
    inputs.update(overrides)

    return build_executive_markdown_report(**inputs)


def docx_report(**overrides):
    inputs = {
        "brand": "Espresso House",
        "market": "Berlin",
        "category": "cafes",
        "audience": "remote workers",
        "summary_df": sample_summary_df(),
        "top_brands_df": sample_top_brands_df(),
        "recommendations": "Prioritize AI-citable category evidence.",
        "strategy_report": "Use future benchmark cycles to validate progress.",
        "gap_analysis": "The tested answers suggest limited prompt-level visibility.",
        "run_mode": "Full Report Mode",
        "prompt_limit": None,
        "brand_intelligence": {
            "recommendation_drivers": "Diagnostic inference from benchmark answers.",
            "target_brand_understanding": "AI-inferred target brand understanding.",
            "positioning_gap_analysis": "Positioning gaps require validation.",
        },
        "prompt_categories": ["Best Options", "Local Recommendations"],
        "geo_content_roadmap": None,
    }
    inputs.update(overrides)

    return create_executive_docx_report(**inputs)


def extract_docx_text(docx_bytes):
    document = Document(BytesIO(docx_bytes))
    text_parts = []

    text_parts.extend(paragraph.text for paragraph in document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text_parts.extend(paragraph.text for paragraph in cell.paragraphs)

    return "\n".join(part for part in text_parts if part)


def test_markdown_quick_test_report_is_not_client_deliverable():
    report = markdown_report(
        run_mode="Quick Test Mode",
        prompt_limit=1,
        deliverable_status="Development-only limited-prompt output. Not client-deliverable.",
    )

    assert_has_phrase_family(report, ["test version only", "quick test mode"])
    assert_has_phrase_family(report, ["not client-deliverable", "not client deliverable"])


def test_markdown_report_preserves_methodology_limitations():
    report = markdown_report()

    assert_has_phrase_family(report, ["ai visibility benchmark"])
    assert_has_phrase_family(report, ["not market share"])
    assert_has_phrase_family(report, ["sales performance report", "business performance outcomes"])
    assert_has_phrase_family(report, ["clinical evaluation"])


def test_markdown_report_uses_benchmark_safe_framing():
    report = markdown_report()

    assert_has_phrase_family(report, ["tested prompt set", "benchmark"])


def test_markdown_report_avoids_forbidden_overclaim_phrases():
    assert_forbidden_overclaims_absent(markdown_report())


def test_markdown_brand_intelligence_keeps_diagnostic_separation():
    report = markdown_report()

    assert_has_phrase_family(
        report,
        [
            "diagnostic",
            "not part of visibility scoring",
            "not included in scoring",
        ],
    )


def test_docx_quick_test_report_is_not_client_deliverable():
    report_text = extract_docx_text(
        docx_report(run_mode="Quick Test Mode", prompt_limit=1)
    )

    assert_has_phrase_family(report_text, ["test version only", "quick test mode"])
    assert_has_phrase_family(report_text, ["not client-deliverable", "not client deliverable"])


def test_docx_report_preserves_methodology_limitations():
    report_text = extract_docx_text(docx_report())

    assert_has_phrase_family(report_text, ["ai visibility benchmark"])
    assert_has_phrase_family(report_text, ["not market share"])
    assert_has_phrase_family(report_text, ["sales performance report", "business performance outcomes"])
    assert_has_phrase_family(report_text, ["clinical evaluation"])


def test_docx_report_avoids_forbidden_overclaim_phrases():
    assert_forbidden_overclaims_absent(extract_docx_text(docx_report()))


def test_docx_brand_intelligence_keeps_diagnostic_separation():
    report_text = extract_docx_text(docx_report())

    assert_has_phrase_family(
        report_text,
        [
            "diagnostic",
            "not part of visibility scoring",
            "not included in scoring",
        ],
    )

def test_markdown_zero_visibility_report_uses_first_detection_language():
    report = zero_visibility_markdown_report()

    assert_has_phrase_family(
        report,
        [
            "not detected",
            "not retrieved",
            "zero visibility",
        ],
    )
    assert_has_phrase_family(
        report,
        [
            "candidate-set inclusion",
            "candidate set",
            "first detection",
        ],
    )
    assert_has_phrase_family(
        report,
        [
            "tested prompt set",
            "benchmark run",
        ],
    )


def test_markdown_zero_visibility_report_identifies_retrieved_brands():
    report = zero_visibility_markdown_report()

    assert "Munich Re" in report
    assert "Swiss Re" in report
    assert_has_phrase_family(
        report,
        [
            "retrieved",
            "visible",
            "competitor",
            "reference",
        ],
    )


def test_markdown_report_includes_validation_or_future_benchmark_language():
    report = markdown_report()

    assert_has_phrase_family(
        report,
        [
            "future benchmark",
            "future benchmark cycles",
            "future runs",
            "validate progress",
            "validated in future",
        ],
    )


def test_markdown_report_preserves_framework_evidence_boundaries():
    report = markdown_report()

    assert_has_phrase_family(
        report,
        [
            "benchmark",
            "tested prompt set",
        ],
    )
    assert_has_phrase_family(
        report,
        [
            "diagnostic",
            "requires validation",
            "validate",
        ],
    )


def test_markdown_zero_visibility_report_avoids_forbidden_overclaim_phrases():
    assert_forbidden_overclaims_absent(zero_visibility_markdown_report())